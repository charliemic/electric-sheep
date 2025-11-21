#!/usr/bin/env python3
"""
Convert Markdown/HTML to DOCX for Google Docs import.

Usage:
    python scripts/convert-to-docx.py input.md output.docx
    python scripts/convert-to-docx.py input.html output.docx

SECURITY NOTE: This script must NEVER automatically prompt, send, or trigger emails
to other users. Such functionality would be intrusive and is explicitly prohibited.
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed")
    print("Install with: pip3 install python-docx --user")
    sys.exit(1)


def html_to_text(html_content):
    """Convert HTML to plain text, preserving structure."""
    # Remove HTML tags but keep content
    text = re.sub(r'<h1[^>]*>(.*?)</h1>', r'\n# \1\n\n', html_content, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'\n## \1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n### \1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)  # Remove remaining HTML tags
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    return text


def markdown_to_docx(md_content, output_path):
    """Convert Markdown content to DOCX with blog-style formatting."""
    doc = Document()
    
    # Web-like page setup: wider page, minimal margins (like a web page)
    from docx.shared import Inches, Cm
    sections = doc.sections
    for section in sections:
        # Web-like: wider page, minimal margins
        section.page_width = Inches(11)  # Wider than standard 8.5"
        section.page_height = Inches(14)  # Taller to accommodate web-like layout
        section.top_margin = Inches(0.3)  # Minimal top margin
        section.bottom_margin = Inches(0.3)  # Minimal bottom margin
        section.left_margin = Inches(1.0)   # Narrow side margins
        section.right_margin = Inches(1.0)
    
    # Blog-style formatting: better typography and spacing
    # Set default paragraph style (blog-like)
    # Using modern, clean fonts that match web/blog style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'  # Modern, clean sans-serif (widely available, web-like)
    font.size = Pt(11)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.15  # Tighter line spacing (more compact)
    paragraph_format.space_after = Pt(3)  # Minimal space between paragraphs
    
    # Style Heading 1 (blog title style - slightly smaller)
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'  # Match body font for consistency
    h1_font.size = Pt(22)  # Slightly smaller, more proportional
    h1_font.bold = True
    h1_format = h1_style.paragraph_format
    h1_format.space_after = Pt(12)  # Less space after title
    h1_format.space_before = Pt(6)  # Less space before title
    
    # Style Heading 2 (section headings)
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(18)
    h2_font.bold = True
    h2_format = h2_style.paragraph_format
    h2_format.space_after = Pt(8)  # Less space after headings
    h2_format.space_before = Pt(16)  # Less space before sections (web-like)
    
    # Style Heading 3 (subsections)
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_format = h3_style.paragraph_format
    h3_format.space_after = Pt(8)
    h3_format.space_before = Pt(16)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Empty line - skip or add minimal spacing (web-like, less whitespace)
            # Only add paragraph if it's not excessive
            if i > 0 and lines[i-1].strip():  # Only if previous line had content
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)  # No extra space
            i += 1
            continue
        
        # Headers (with blog-style spacing)
        if line.startswith('# '):
            # Add space before main heading
            if i > 0:
                doc.add_paragraph()
            p = doc.add_heading(line[2:], level=1)
            i += 1
        elif line.startswith('## '):
            # Add space before section headings
            doc.add_paragraph()
            p = doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            i += 1
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        elif re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s+', '', line), style='List Number')
            i += 1
        # Code blocks
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            i += 1
        # Regular paragraph
        else:
            # Process inline formatting
            para_text = line
            p = doc.add_paragraph()
            
            # Handle bold, italic, code
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', para_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'Calibri'  # Bold text matches body font
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    # Add background color for code (light gray) - simplified approach
                    # Note: Background color in DOCX is complex, Google Docs may not fully support it
                elif part:
                    run = p.add_run(part)
                    run.font.name = 'Calibri'  # Ensure body text uses consistent font
            
            i += 1
    
    doc.save(output_path)
    print(f"✓ DOCX file created: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python scripts/convert-to-docx.py <input.md|input.html> <output.docx>")
        sys.exit(1)
    
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)
    
    # Read input file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Convert HTML to markdown if needed
    if input_path.suffix.lower() == '.html':
        content = html_to_text(content)
    
    # Convert to DOCX
    markdown_to_docx(content, output_path)


if __name__ == '__main__':
    main()

